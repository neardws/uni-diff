import os
from typing import List
from .base import BaseConverter, ConvertedDocument, TextBlock


class DOCXConverter(BaseConverter):
    """Converter for Microsoft Word DOCX files."""

    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx']

    def convert(self, file_path: str) -> ConvertedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        blocks = []
        full_text = ""

        try:
            from docx import Document
            doc = Document(file_path)

            y_offset = 0
            for para in doc.paragraphs:
                text = para.text
                if text.strip():
                    full_text += text + "\n"
                    blocks.append(TextBlock(
                        text=text,
                        page=0,
                        x=0,
                        y=y_offset,
                        width=len(text) * 7,
                        height=14,
                        metadata={'style': para.style.name if para.style else 'Normal'}
                    ))
                y_offset += 14

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text = " | ".join(row_text)
                        full_text += text + "\n"
                        blocks.append(TextBlock(
                            text=text,
                            page=0,
                            x=0,
                            y=y_offset,
                            width=len(text) * 7,
                            height=14,
                            metadata={'type': 'table_row'}
                        ))
                    y_offset += 14

        except ImportError:
            import subprocess
            try:
                result = subprocess.run(
                    ['pandoc', '-t', 'plain', file_path],
                    capture_output=True,
                    text=True,
                    check=True
                )
                full_text = result.stdout
                lines = full_text.split('\n')
                y_offset = 0
                for line in lines:
                    if line.strip():
                        blocks.append(TextBlock(
                            text=line,
                            page=0,
                            x=0,
                            y=y_offset,
                            width=len(line) * 7,
                            height=12
                        ))
                    y_offset += 12
            except FileNotFoundError:
                raise RuntimeError(
                    "Neither python-docx nor pandoc is available. "
                    "Install python-docx or pandoc."
                )

        return ConvertedDocument(
            blocks=blocks,
            full_text=full_text,
            page_count=1,
            source_path=file_path,
            source_type='docx'
        )
